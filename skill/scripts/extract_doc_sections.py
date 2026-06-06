#!/usr/bin/env python3
"""
extract_doc_sections.py — Document Section Extractor
Part of the evaluation-framework skill.

Extracts structured information from PDF, DOCX, and TXT files.

Usage:
    python extract_doc_sections.py --file path/to/document.pdf
    python extract_doc_sections.py --file path/to/document.docx
    python extract_doc_sections.py --file path/to/document.txt
    python extract_doc_sections.py --file path/to/document.md

Output (JSON):
    {
        "file_name": "project_brief.pdf",
        "file_type": "pdf",
        "page_count": 12,
        "word_count": 3400,
        "extracted_sections": {
            "objective": "...",
            "scope": "...",
            "stakeholders": "...",
            "timeline": "...",
            "budget": "...",
            "success_criteria": "...",
            "risks": "...",
            "methodology": "...",
            "team": "...",
            "context": "..."
        },
        "section_status": {
            "objective": "found",
            "scope": "partial",
            ...
        },
        "raw_headings": ["Executive Summary", "Project Scope", ...],
        "evidence_map": [
            {"section": "risks", "text": "...", "location": "Page 4, Section 'Risk Register'"}
        ],
        "warnings": []
    }
"""

import sys
import json
import argparse
import re
from pathlib import Path


# ─── Section Target Definitions ────────────────────────────────────────────────

SECTION_PATTERNS = {
    "objective": {
        "heading_patterns": [
            r"executive summary", r"overview", r"introduction", r"abstract",
            r"project objective", r"goal", r"aim", r"purpose", r"mission",
            r"background", r"problem statement", r"business case",
        ],
        "content_patterns": [
            r"the (goal|objective|purpose|aim) (of this|is to|of the)",
            r"we (aim|seek|intend) to",
            r"this (project|initiative|program) (will|aims|seeks)",
            r"the (project|initiative) is designed to",
        ],
    },
    "scope": {
        "heading_patterns": [
            r"scope", r"project scope", r"in scope", r"out of scope",
            r"boundaries", r"coverage", r"included", r"excluded",
            r"deliverables",
        ],
        "content_patterns": [
            r"in scope", r"out of scope", r"this (project|initiative) (includes|covers|excludes)",
            r"the following (are|is) (included|excluded|in scope|out of scope)",
        ],
    },
    "stakeholders": {
        "heading_patterns": [
            r"stakeholder", r"key (stakeholder|contact|player)",
            r"project team", r"governance", r"accountability",
            r"raci", r"responsibility", r"sponsor",
        ],
        "content_patterns": [
            r"(key |primary |main )stakeholder",
            r"(project |executive |senior )(sponsor|owner|champion)",
            r"(responsible|accountable) (for|to)",
        ],
    },
    "timeline": {
        "heading_patterns": [
            r"timeline", r"schedule", r"gantt", r"roadmap", r"milestone",
            r"phase", r"project plan", r"delivery plan", r"implementation plan",
        ],
        "content_patterns": [
            r"(q[1-4]|quarter \d|fy\d{2,4})",
            r"(by|before|until|no later than).{0,20}(january|february|march|april|may|june|july|august|september|october|november|december|\d{4})",
            r"(week|month|sprint) \d",
            r"\d+ (weeks?|months?|days?) (from|after|following)",
        ],
    },
    "budget": {
        "heading_patterns": [
            r"budget", r"cost", r"financial", r"investment", r"funding",
            r"resource allocation", r"financial plan", r"cost estimate",
        ],
        "content_patterns": [
            r"\$[\d,]+(?:\.\d+)?(?:\s*(?:million|billion|k|m|b))?",
            r"€[\d,]+", r"£[\d,]+",
            r"(budget|cost|investment|funding).{0,30}(million|billion|thousand|\$|€|£)",
            r"(allocated|approved|requested).{0,20}(budget|funding|amount)",
        ],
    },
    "success_criteria": {
        "heading_patterns": [
            r"success criteria", r"success factor", r"kpi", r"key (performance|result)",
            r"metric", r"target", r"goal", r"objective",
            r"acceptance criteria", r"definition of (done|success)",
            r"okr", r"measure of success",
        ],
        "content_patterns": [
            r"(success|successful) (is defined|will be measured|criteria)",
            r"(key|primary|core) (metric|kpi|indicator)",
            r"(achieve|reach|attain).{0,30}(target|goal|metric)",
            r"measured by", r"tracked (via|through|by)",
        ],
    },
    "risks": {
        "heading_patterns": [
            r"risk", r"assumption", r"dependency", r"constraint", r"issue",
            r"challenge", r"concern", r"limitation", r"caveat",
            r"mitigation", r"contingency",
        ],
        "content_patterns": [
            r"(key|primary|main|critical) (risk|assumption|dependency)",
            r"(high|medium|low) (risk|likelihood|impact|probability)",
            r"(mitigat|contingenc|address)",
            r"risk.{0,20}(is|are|include|may|could)",
        ],
    },
    "methodology": {
        "heading_patterns": [
            r"methodology", r"approach", r"method", r"how we (will|plan)",
            r"implementation approach", r"process", r"framework",
            r"design approach", r"technical approach",
        ],
        "content_patterns": [
            r"(our|the) (approach|methodology|process|method) (is|will be|involves)",
            r"(agile|waterfall|scrum|kanban|lean|design thinking)",
            r"(phase \d|step \d|stage \d).{0,50}(will|involves|includes)",
        ],
    },
    "team": {
        "heading_patterns": [
            r"team", r"resource", r"staff", r"headcount", r"organization",
            r"who (will|is|are)", r"project (team|structure|organization)",
            r"people", r"talent",
        ],
        "content_patterns": [
            r"(project manager|product manager|tech lead|architect|analyst|developer)",
            r"(full.?time|part.?time|fte|contractor)",
            r"\d+\s*(people|person|staff|fte|team member|resource)",
        ],
    },
    "context": {
        "heading_patterns": [
            r"background", r"context", r"situation", r"current state",
            r"problem", r"challenge", r"why (this|we|now)",
            r"market context", r"industry context",
        ],
        "content_patterns": [
            r"(currently|at present|today).{0,30}(we|the (market|company|organization|system))",
            r"(the (problem|challenge|issue|opportunity) (is|that))",
            r"(growing|declining|changing|evolving).{0,30}(market|industry|demand|trend)",
        ],
    },
}


# ─── Text Extraction Functions ──────────────────────────────────────────────────

def extract_from_txt(file_path: Path) -> tuple[str, dict]:
    """Extract text from plain text or markdown file."""
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    metadata = {
        "page_count": text.count("\n\n"),  # Rough paragraph count
        "word_count": len(text.split()),
    }
    return text, metadata


def extract_from_pdf(file_path: Path) -> tuple[str, dict]:
    """Extract text from PDF. Falls back gracefully if pypdf not available."""
    try:
        import pypdf
        reader = pypdf.PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        full_text = "\n\n".join(pages)
        metadata = {
            "page_count": len(reader.pages),
            "word_count": len(full_text.split()),
        }
        return full_text, metadata
    except ImportError:
        return (
            f"[PDF_EXTRACTION_UNAVAILABLE: pypdf not installed. "
            f"Install with: pip install pypdf --break-system-packages]",
            {"page_count": 0, "word_count": 0},
        )
    except Exception as e:
        return f"[PDF_EXTRACTION_ERROR: {str(e)}]", {"page_count": 0, "word_count": 0}


def extract_from_docx(file_path: Path) -> tuple[str, dict]:
    """Extract text from DOCX. Falls back gracefully if python-docx not available."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Prefix headings for better section detection
                if para.style.name.startswith("Heading"):
                    paragraphs.append(f"\n## {para.text.strip()}\n")
                else:
                    paragraphs.append(para.text.strip())
        full_text = "\n".join(paragraphs)
        metadata = {
            "page_count": len(doc.paragraphs) // 25,  # Rough estimate
            "word_count": len(full_text.split()),
        }
        return full_text, metadata
    except ImportError:
        return (
            f"[DOCX_EXTRACTION_UNAVAILABLE: python-docx not installed. "
            f"Install with: pip install python-docx --break-system-packages]",
            {"page_count": 0, "word_count": 0},
        )
    except Exception as e:
        return f"[DOCX_EXTRACTION_ERROR: {str(e)}]", {"page_count": 0, "word_count": 0}


# ─── Section Extraction ────────────────────────────────────────────────────────

def split_into_sections(text: str) -> list[dict]:
    """Split document text into heading + content blocks."""
    sections = []
    
    # Split on markdown-style headings or ALL-CAPS lines (common in PDFs)
    pattern = re.compile(
        r"(#{1,4}\s+.+|^[A-Z][A-Z\s]{5,50}$)",
        re.MULTILINE
    )
    
    matches = list(pattern.finditer(text))
    
    if not matches:
        # No clear headings — treat whole doc as one block
        return [{"heading": "FULL_DOCUMENT", "content": text[:8000], "start": 0}]
    
    for i, match in enumerate(matches):
        heading = match.group().strip("# ").strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()
        sections.append({
            "heading": heading,
            "content": content[:2000],  # Cap at 2000 chars per section
            "start": match.start(),
        })
    
    return sections


def find_section_match(sections: list[dict], target: str, patterns: dict) -> tuple[str, str]:
    """
    Find the section in the document that best matches a target (e.g., 'objective').
    Returns (extracted_text, status) where status is 'found', 'partial', or 'not_found'.
    """
    heading_patterns = patterns.get("heading_patterns", [])
    content_patterns = patterns.get("content_patterns", [])
    
    best_match = None
    best_score = 0
    
    for section in sections:
        score = 0
        heading_lower = section["heading"].lower()
        content_lower = section["content"].lower()
        
        # Score heading matches (higher weight)
        for hp in heading_patterns:
            if re.search(hp, heading_lower):
                score += 10
                break
        
        # Score content pattern matches
        for cp in content_patterns:
            if re.search(cp, content_lower, re.IGNORECASE):
                score += 3
        
        if score > best_score:
            best_score = score
            best_match = section
    
    if best_score >= 10:
        return best_match["content"][:1500], "found"
    elif best_score >= 3:
        return best_match["content"][:1000], "partial"
    else:
        return "", "not_found"


# ─── Evidence Map Builder ──────────────────────────────────────────────────────

def build_evidence_map(sections: list[dict], section_status: dict) -> list[dict]:
    """Build a list of evidence references for the report appendix."""
    evidence = []
    for target, status in section_status.items():
        if status in ("found", "partial"):
            for section in sections:
                # Find which section contributed
                heading_patterns = SECTION_PATTERNS[target]["heading_patterns"]
                for hp in heading_patterns:
                    if re.search(hp, section["heading"].lower()):
                        snippet = section["content"][:200].replace("\n", " ").strip()
                        if snippet:
                            evidence.append({
                                "section": target,
                                "text": snippet + ("..." if len(section["content"]) > 200 else ""),
                                "location": f"Section: '{section['heading']}'",
                            })
                        break
    return evidence


# ─── Main Extraction Orchestrator ──────────────────────────────────────────────

def extract_document(file_path: Path) -> dict:
    """Main function: extract and structure all information from a document."""
    file_type = file_path.suffix.lower().lstrip(".")
    
    # Dispatch to correct extractor
    if file_type == "pdf":
        raw_text, metadata = extract_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        raw_text, metadata = extract_from_docx(file_path)
    elif file_type in ("txt", "md", "rst"):
        raw_text, metadata = extract_from_txt(file_path)
    else:
        return {
            "error": f"Unsupported file type: .{file_type}",
            "supported_types": ["pdf", "docx", "doc", "txt", "md", "rst"],
        }
    
    # Split into sections
    sections = split_into_sections(raw_text)
    raw_headings = [s["heading"] for s in sections if s["heading"] != "FULL_DOCUMENT"]
    
    # Extract each target section
    extracted_sections = {}
    section_status = {}
    warnings = []
    
    for target, patterns in SECTION_PATTERNS.items():
        text, status = find_section_match(sections, target, patterns)
        extracted_sections[target] = text
        section_status[target] = status
    
    # Build evidence map
    evidence_map = build_evidence_map(sections, section_status)
    
    # Warnings
    found_count = sum(1 for s in section_status.values() if s == "found")
    partial_count = sum(1 for s in section_status.values() if s == "partial")
    
    if found_count + partial_count < 3:
        warnings.append(
            "Fewer than 3 sections could be extracted. "
            "Document may be image-based (scanned PDF) or have non-standard formatting. "
            "Consider entering Clarification Mode."
        )
    
    if metadata.get("word_count", 0) > 15000:
        warnings.append(
            f"Document is long ({metadata['word_count']} words). "
            "Only key sections have been extracted. Some detail may be missed."
        )
    
    return {
        "file_name": file_path.name,
        "file_type": file_type,
        "page_count": metadata.get("page_count", 0),
        "word_count": metadata.get("word_count", 0),
        "extracted_sections": extracted_sections,
        "section_status": section_status,
        "raw_headings": raw_headings[:20],  # First 20 headings
        "evidence_map": evidence_map,
        "warnings": warnings,
        "summary": {
            "sections_found": found_count,
            "sections_partial": partial_count,
            "sections_missing": len(SECTION_PATTERNS) - found_count - partial_count,
            "quality": "good" if found_count >= 5 else ("acceptable" if found_count >= 3 else "poor"),
        },
    }


# ─── CLI Entry Point ────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Extract structured sections from project documents (PDF, DOCX, TXT)."
    )
    parser.add_argument("--file", required=True, type=str, help="Path to document file")
    parser.add_argument("--pretty", action="store_true", help="Pretty-print JSON output")
    parser.add_argument(
        "--section", type=str, help="Extract only a specific section (e.g., 'risks')"
    )
    
    args = parser.parse_args()
    file_path = Path(args.file)
    
    if not file_path.exists():
        print(json.dumps({"error": f"File not found: {args.file}"}))
        sys.exit(1)
    
    result = extract_document(file_path)
    
    # If specific section requested, return just that
    if args.section and args.section in result.get("extracted_sections", {}):
        section_result = {
            "section": args.section,
            "content": result["extracted_sections"][args.section],
            "status": result["section_status"][args.section],
        }
        print(json.dumps(section_result, indent=2 if args.pretty else None))
    else:
        print(json.dumps(result, indent=2 if args.pretty else None))


if __name__ == "__main__":
    main()
